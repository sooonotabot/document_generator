from flask import Flask, render_template, request, redirect, url_for
from docx import Document
import os
from werkzeug.utils import secure_filename


app = Flask(__name__)
UPLOAD_FOLDER = 'generated_docs'
app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(), UPLOAD_FOLDER)
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        name = request.form['user_name']
        title = request.form['title']
        document = Document()
        document.add_heading('Generated Document', 0)
        document.add_paragraph(f"Name: {name}")
        filename = f"{title}.docx"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        document.save(file_path)
        return redirect(url_for('success', filename=filename))
    return render_template('form.html')

@app.route('/success/<filename>')
def success(filename):
    return f"Document '{filename}' generated and saved successfully"

if __name__ == '__main__':
    app.run(debug=True)
